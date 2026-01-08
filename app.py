import os
import json
import re
from typing import List, Any, Union, Dict
import streamlit as st
import datetime
from docx import Document
from docx.shared import Inches
from io import BytesIO
import nest_asyncio
nest_asyncio.apply()  # <--- The Magic Line. Fixes the loop issues.
from agent_runner import (
    run_understanding_agent,
    run_mapping_agent,
    run_product_selector_agent,
    run_feedback_agent,
    run_final_output_agent
)

# --- MOCK DATA CONSTANTS ---
MOCK_UNDERSTANDING_JSON = {
    "process_map": [
        {"step_name": "Employee submits expense report", "description": "Fills Excel template.", "role": "Employee", "decision_point": False},
        {"step_name": "Manager reviews report", "description": "Checks for receipts.", "role": "Manager", "decision_point": True, "condition": "Total < $500?"},
        {"step_name": "Finance final approval", "description": "Checks tax compliance.", "role": "Finance", "decision_point": False},
        {"step_name": "Payment processing", "description": "Manual entry to portal.", "role": "Finance", "decision_point": False}
    ],
    "issues": ["Manual entry errors", "Lost receipts"],
    "opportunities": ["Automate banking transfer", "Digital receipt capture"]
}

MOCK_PRODUCT_RESULT = {
    "top_5_tools": ["Expensify", "Power Automate", "AppSheet", "Rydoo", "SAP Concur"],
    "recommended_tool": "Power Automate",
    "reason_for_recommendation": "Strong fit for Microsoft ecosystem and simple approval flows."
}

# --- STREAMLIT CONFIG ---
st.set_page_config(page_title="Excel Process Mapping", layout="wide")



if "saved_reports" not in st.session_state:
    st.session_state["saved_reports"] = []  # List of {timestamp, content, png_bytes, title}

if "is_dev" not in st.session_state:
    st.session_state["is_dev"] = False

# --- HELPER FUNCTIONS ---

def _clear_analysis_state():
    """Wipes the 'Work in Progress' data to reset for a new file."""
    keys_to_clear = [
        "understanding_json", 
        "map_result", 
        "product_result", 
        "feedback_result", 
        "feedback_history", 
        "feedback_log"
    ]
    for key in keys_to_clear:
        if key in st.session_state:
            del st.session_state[key]

def _save_report_to_app(report_content, map_path):
    """Saves the current report to the sidebar list."""
    timestamp = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")
    
    # Extract a title from the report or use default
    title = f"Report - {timestamp}"
    
    # Store PNG bytes so we don't rely on the file system for saved reports
    png_bytes = None
    if map_path and os.path.exists(map_path):
        with open(map_path, "rb") as f:
            png_bytes = f.read()

    report_entry = {
        "id": len(st.session_state["saved_reports"]) + 1,
        "title": title,
        "content": report_content,
        "png_bytes": png_bytes,
        "timestamp": timestamp
    }
    st.session_state["saved_reports"].append(report_entry)
    st.toast(f"✅ Saved '{title}' to App!")

st.title("📊 Excel Process Mapping Assistant")

def _extract_png_bytes(map_result: Any) -> bytes | None:
    if isinstance(map_result, str) and os.path.exists(map_result):
        with open(map_result, "rb") as f:
            return f.read()
    elif isinstance(map_result, (bytes, bytearray)):
        return bytes(map_result)
    return None

def _generate_docx(report_text: str, image_bytes: bytes = None, image_path: str = None) -> BytesIO:
    """Converts Markdown to Docx (handles both file path and raw bytes for images)."""
    doc = Document()
    doc.add_heading('Final Process Report', 0)
    lines = report_text.split('\n')
    
    for line in lines:
        line = line.strip()
        if line.startswith('# '): doc.add_heading(line[2:], level=1)
        elif line.startswith('## '): doc.add_heading(line[3:], level=2)
        elif line.startswith('### '): doc.add_heading(line[4:], level=3)
        elif "![Process Map]" in line:
            doc.add_heading('Process Flowchart', level=2)
            try:
                # Prioritize bytes (for saved reports), fall back to path (for live)
                if image_bytes:
                    doc.add_picture(BytesIO(image_bytes), width=Inches(6))
                elif image_path and os.path.exists(image_path):
                    doc.add_picture(image_path, width=Inches(6))
                else:
                    doc.add_paragraph("[Image Missing]")
            except Exception as e:
                doc.add_paragraph(f"[Error rendering image: {e}]")
        elif line.startswith('* ') or line.startswith('- '):
            doc.add_paragraph(line[2:].replace('**', ''), style='List Bullet')
        elif line:
            doc.add_paragraph(line.replace('**', ''))
            
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def _smart_parse_json(data: Any) -> Dict:
    """Robustly parses JSON from LLM output, handling Markdown code blocks."""
    if isinstance(data, dict):
        return data
        
    if isinstance(data, str):
        cleaned_text = re.sub(r"```json\s*|\s*```", "", data, flags=re.IGNORECASE).strip()
        try:
            return json.loads(cleaned_text)
        except json.JSONDecodeError:
            try:
                start = cleaned_text.find('{')
                end = cleaned_text.rfind('}') + 1
                if start != -1 and end != -1:
                    return json.loads(cleaned_text[start:end])
            except:
                pass
        return {"raw_text": data, "error": "Could not parse JSON"}
    return {}

def _render_formatted_understanding(data: Dict):
    """Renders the Understanding JSON as nicely formatted text."""
    
    # 1. Process Map Section
    st.markdown("# 🗺️ Process Map")
    steps = data.get("process_map", [])
    
    if not steps:
        st.info("No process steps found.")
    else:
        for i, step in enumerate(steps, 1):
            name = step.get("step_name", "Untitled Step")
            role = step.get("role", "Unknown Role")
            desc = step.get("description", "")
            
            # Formatting as requested: H2 for Step, Body for details
            st.markdown(f"### {i}. {name}")
            st.markdown(f"**Role:** {role}")
            st.markdown(f"{desc}")

    # 2. Issues Section
    st.markdown("# ⚠️ Issues")
    issues = data.get("issues", [])
    if issues:
        for issue in issues:
            st.markdown(f"- {issue}")
    else:
        st.success("No issues identified.")

    # 3. Opportunities Section
    st.markdown("# 💡 Opportunities")
    opps = data.get("opportunities", [])
    if opps:
        for opp in opps:
            st.markdown(f"- {opp}")
    else:
        st.info("No opportunities identified.")

def _render_formatted_product_recs(data: Dict):
    """Renders Product Recommendations with visual hierarchy."""
    
    
    rec = data.get("recommended_tool")
    reason = data.get("reason_for_recommendation")
    
    if rec and rec != "null" and rec is not None:
        st.markdown(f"### 🏆 Recommended Tool: **{rec}**")
        if reason:
            st.markdown("#### Reason for Recommendation")
            st.markdown(reason)
    else:
        st.warning("No specific tool recommended.")

    st.markdown("---")
    st.markdown("### 🥈 Top 5 Contenders")
    
    tools = data.get("top_5_tools", [])
    if tools:
        for i, tool in enumerate(tools, 1):
            st.markdown(f"{i}. **{tool}**")
    else:
        st.caption("No other contenders found.")

def _mark_as_saved():
    """Callback to mark the current report as 'safe'."""
    st.session_state["report_is_saved"] = True

# --- MODALS (DIALOGS) ---

@st.dialog("💾 Unsaved Report Detected")
def show_unsaved_warning(current_report, map_path):
    st.warning("You have a generated report on screen. Starting a new analysis will lose it unless you save.")
    
    col1, col2, col3 = st.columns(3)
    
    with col1:
        if st.button("📂 Save in App", use_container_width=True):
            _save_report_to_app(current_report, map_path)
            # Clear the 'current' report so we can proceed
            del st.session_state["final_output"]
            st.rerun()
            
    with col2:
        # Generate DOCX for instant download
        docx = _generate_docx(current_report, image_path=map_path)
        st.download_button("📄 Save as Word", data=docx, file_name="Report.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
        
    with col3:
        if st.button("🗑️ Discard & Run", type="primary", use_container_width=True):
            del st.session_state["final_output"]
            st.rerun()

@st.dialog("📄 Saved Report Viewer", width="large")
def view_saved_report(report):
    st.markdown(f"### {report['title']}")
    
    # 1. Render content (handle image replacement)
    parts = report["content"].split("![Process Map](process_map.png)")
    st.markdown(parts[0])
    if report["png_bytes"]:
        st.image(report["png_bytes"], caption="Process Snapshot")
    st.markdown(parts[1] if len(parts) > 1 else "")
    
    st.divider()
    
    # 2. Download Options
    c1, c2 = st.columns(2)
    with c1:
        st.download_button("📥 Download Markdown", report["content"], file_name=f"{report['title']}.md")
    with c2:
        docx = _generate_docx(report["content"], image_bytes=report["png_bytes"])
        st.download_button("📥 Download Word Doc", docx, file_name=f"{report['title']}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

@st.dialog("🔐 Developer Login")
def show_dev_login():
    password = st.text_input("Enter Access Key", type="password")
    if st.button("Unlock"):
        if password == "2026admin123":
            st.session_state["is_dev"] = True
            st.rerun()
        else:
            st.error("Incorrect password.")
# --- MAIN UI ---
# --- SIDEBAR & CONFIG ---
with st.sidebar:
    st.header("Workflow")
    st.markdown("1️⃣ Upload Excel files")
    st.markdown("2️⃣ Click **Run analysis**")
    st.markdown("3️⃣ Review & Feedback")
    st.markdown("4️⃣ Final Output")
    
    st.header("🗂️ Saved Reports")
    if not st.session_state["saved_reports"]:
        st.caption("No reports saved yet.")
    else:
        for rep in st.session_state["saved_reports"]:
            # Use a button to trigger the dialog
            if st.button(f"📄 {rep['title']}", key=f"btn_{rep['id']}"):
                view_saved_report(rep)

    st.markdown("---")
    # --- CONDITIONAL DEBUG SECTION ---
    if st.session_state["is_dev"]:
        
        
        st.subheader("🛠️ Debug Settings")
        use_mock_understanding = st.checkbox("Mock Step 2 (Understanding)", value=True)
        use_mock_mapping = st.checkbox("Mock Step 3a (Mapping)", value=False)
        use_mock_product = st.checkbox("Mock Step 3b (Product Selector)", value=False)
        
        st.success("👨‍💻 Developer Mode Active")

        if st.button("Logout"):
            st.session_state["is_dev"] = False
            st.rerun()
            
    else:
        # Defaults for normal users (Must be False so the real agents run!)
        use_mock_understanding = False
        use_mock_mapping = False
        use_mock_product = False
        
        # The "Secret Door" Button
        if st.button("🔒 Developer Access", type="tertiary"):
            show_dev_login()

    

st.markdown("""
**Step 1.** Upload Excel files.
**Step 2.** Click **“🚀 Run analysis”**.
**Step 3.** Review outputs & provide feedback.
""")

# STEP 1: UPLOAD
uploaded_files = st.file_uploader("Upload 📂 Excel files", type=["xlsx", "xls", "xlsm"], accept_multiple_files=True)

if uploaded_files and len(uploaded_files) > 3:
    st.warning("Only first 3 files will be used.")
    uploaded_files = uploaded_files[:3]

if uploaded_files:
    st.success(f"{len(uploaded_files)} file(s) ready.")

# STEP 2 & 3: RUN ANALYSIS TRIGGER
run_button = st.button("🚀 Run analysis", type="primary")

if run_button:
   # INTERCEPTION: Check if there is an unsaved report on screen
    if "final_output" in st.session_state and not st.session_state.get("report_is_saved", False):
        show_unsaved_warning(
            st.session_state["final_output"], 
            "process_map.png" # Assuming this is the current live file
        )
    else:
        # PROCEED: No unsaved report, safe to run

        _clear_analysis_state()
        if "final_output" in st.session_state:
            del st.session_state["final_output"]
        if "report_is_saved" in st.session_state:
            del st.session_state["report_is_saved"]

        if not uploaded_files and not use_mock_understanding:
            st.error("Please upload a file or enable Mock Understanding.")
            st.stop()

    # --- STEP 2: UNDERSTANDING AGENT ---
    if use_mock_understanding:
        st.warning("⚠️ Using MOCK DATA for Understanding Agent.")
        understanding_json = MOCK_UNDERSTANDING_JSON
    else:
        with st.spinner("Running Understanding Agent (Gemini)..."):
            raw_understanding = run_understanding_agent(uploaded_files)
            understanding_json = _smart_parse_json(raw_understanding)
    
    st.session_state["understanding_json"] = understanding_json

    # --- STEP 3: MAPPING & PRODUCT ---
    
    # 3a. Mapping Agent
    if use_mock_mapping:
        st.warning("⚠️ Skipping Mapping Agent (Mock Mode).")
        map_result = 'process_map.png' 
    else:
        with st.spinner("Running Mapping Agent..."):
            map_result = run_mapping_agent(understanding_json)
    
    st.session_state["map_result"] = map_result

    # 3b. Product Selector
    if use_mock_product:
        st.warning("⚠️ Using MOCK DATA for Product Selector.")
        product_result = MOCK_PRODUCT_RESULT
    else:
        with st.spinner("Running Product Selector Agent..."):
            raw_product = run_product_selector_agent(understanding_json)
            product_result = _smart_parse_json(raw_product)
            
    st.session_state["product_result"] = product_result

# --- DISPLAY OUTPUTS ---

# DISPLAY STEP 2: UNDERSTANDING (NEW FORMATTED VIEW)
if "understanding_json" in st.session_state and "final_output" not in st.session_state:
    st.markdown("---")
    st.subheader("Step 2 – Analysis Results")
    
    data = _smart_parse_json(st.session_state["understanding_json"])
    
    # Render the pretty text version
    _render_formatted_understanding(data)
    
    # Optional: Keep raw JSON in a collapsed box for debugging
    with st.expander("🔍 View Raw JSON Data"):
        st.json(data)

# DISPLAY STEP 3: MAP & PRODUCT
if "map_result" in st.session_state or "product_result" in st.session_state:
    st.markdown("---")
    st.subheader("Step 3 – Mapping & Product Recommendations")
    
    col_map, col_prod = st.columns(2)
    
    # Left Column: Map
    with col_map:
        st.markdown("### 🗺️ Process Map")
        map_result = st.session_state.get("map_result")
        
        if map_result:
            png_bytes = _extract_png_bytes(map_result)
            if png_bytes:
                st.image(png_bytes, caption="Generated Process Map", use_container_width=True)
                st.download_button("📥 Download Map", data=png_bytes, file_name="process_map.png", mime="image/png")
            else:
                st.error(f"Map file not found at path: {map_result}")
        else:
            st.info("Mapping Agent was mocked/skipped.")

    # Right Column: Product
    with col_prod:
        prod_result = _smart_parse_json(st.session_state.get("product_result"))
        
        if prod_result:
            # Call the new visual formatter
            _render_formatted_product_recs(prod_result)
            
            # Keep raw JSON hidden but accessible
            with st.expander("🔍 View Raw Product JSON"):
                st.json(prod_result)
        else:
            st.info("Product Agent was mocked/skipped.")

# STEP 4: FEEDBACK LOOP
if "understanding_json" in st.session_state and "map_result" in st.session_state and "final_output" not in st.session_state:
    st.markdown("---")
    st.subheader("Step 4 – Feedback")
    
    if "last_agent_response" not in st.session_state:
        st.session_state["last_agent_response"] = "I'm ready to refine the process. Just let me know what needs changing!"

    # 2. DISPLAY AGENT RESPONSE (Always Visible)
    with st.chat_message("assistant"):
        st.write(st.session_state["last_agent_response"])

    # 2. Feedback Input
    feedback_text = st.text_area("Your feedback", placeholder="e.g. 'Add a step for manager approval...'")
    
    # 1. Display Change Log
    if "feedback_log" in st.session_state:
        # 1. Display Change History (All Changes)
        if "feedback_history" in st.session_state and st.session_state["feedback_history"]:
            with st.expander(f"📝 Change Log ({len(st.session_state['feedback_history'])} updates)", expanded=True):
                
                # Iterate through the list in reverse order (newest first)
                for i, log in enumerate(reversed(st.session_state["feedback_history"]), 1):
                    
                    # Create a mini header for each change
                    st.markdown(f"**Update #{len(st.session_state['feedback_history']) - i + 1}**")
                    
                    # Show the change details
                    st.info(f"Changed: {log.get('changes_made', 'N/A')}")
                    
                    # Optional: Show reason if it exists
                    if log.get("reason_for_update"):
                        st.caption(f"Reason: {log.get('reason_for_update')}")
                    
                    # Add a divider between entries (except the last one)
                    if i < len(st.session_state["feedback_history"]):
                        st.divider()
                        
        elif "feedback_log" in st.session_state:
            # Fallback if history is empty but log exists (rare legacy case)
            with st.expander("📝 Latest Changes Applied", expanded=True):
                st.success(f"Changes: {st.session_state['feedback_log'].get('changes_made')}")

    if "feedback_history" not in st.session_state:
        st.session_state["feedback_history"] = []

    if st.button("✉️ Send Feedback"):
        if not feedback_text.strip():
            st.warning("Please enter feedback.")
        else:
            with st.spinner("Running Feedback Agent..."):
                # Run Agent
                result_str = run_feedback_agent(
                    understanding_json=st.session_state["understanding_json"],
                    product_selection=st.session_state.get("product_result", {}),
                    user_feedback=feedback_text.strip()
                )
                
                # --- UPDATE STATE & RERUN ---
                try:
                    new_data = _smart_parse_json(result_str)
                    
                    # 1. CAPTURE THE VOICE MESSAGE
                    if new_data.get("agent_response_message"):
                        st.session_state["last_agent_response"] = new_data["agent_response_message"]
                    else:
                        # Fallback if agent is silent
                        st.session_state["last_agent_response"] = "Updates applied successfully."

                    # A. Update Map Path
                    if new_data.get("updated_process_diagram_path"):
                        st.session_state["map_result"] = new_data["updated_process_diagram_path"]
                        
                    # B. Update Product Recs
                    if new_data.get("updated_recommended_tool"):
                        current_prods = _smart_parse_json(st.session_state.get("product_result", {}))
                        current_prods["recommended_tool"] = new_data["updated_recommended_tool"]
                        current_prods["reason_for_recommendation"] = new_data.get("updated_reason_for_tool", "")
                        st.session_state["product_result"] = current_prods
                            
                    # C. Update Core JSON
                    if new_data.get("updated_process_map"):
                        st.session_state["understanding_json"]["process_map"] = new_data["updated_process_map"]

                    # D. Save Log & Refresh
                    # Only save to history if actual changes occurred.
                    changes_text = new_data.get("changes_made", "").lower()
                    
                    # List of phrases that mean "Nothing happened"
                    no_change_phrases = ["none", "no changes", "no changes made", "n/a"]
                    
                    is_meaningful_change = (
                        changes_text and 
                        changes_text not in no_change_phrases and 
                        "no changes" not in changes_text
                    )

                    if is_meaningful_change:
                        st.session_state["feedback_log"] = new_data
                        st.session_state["feedback_history"].append(new_data) # Only append real work
                        print("DEBUG: Change logged.")
                    else:
                        print("DEBUG: Skipped logging (Conversation only).")
                    st.rerun() 
                    
                except Exception as e:
                    st.error(f"Failed: {e}")


# STEP 5: FINAL OUTPUT (Updated)
# ... (keep your existing imports and setup) ...

# --- HELPER TO RENDER MIXED CONTENT ---
def _render_report_with_image(report_text):
    """
    Splits the markdown report to insert the actual Streamlit image widget
    where the placeholder text exists.
    """
    marker = "![Process Map](process_map.png)"
    
    if marker in report_text:
        parts = report_text.split(marker)
        
        # 1. Render text BEFORE the image
        st.markdown(parts[0]) 
        
        # 2. Render the actual image from disk
        if os.path.exists("process_map.png"):
            st.image("process_map.png", caption="Process Flowchart")
        else:
            st.warning("⚠️ Process Map image file not found on disk.")
            
        # 3. Render text AFTER the image
        if len(parts) > 1:
            st.markdown(parts[1])
    else:
        # If the agent didn't output the tag exactly, just show the text
        st.markdown(report_text)


# STEP 5: FINAL OUTPUT (Updated)
if "understanding_json" in st.session_state:
    st.markdown("---")
    st.subheader("Step 5 – Final Deliverable")
    
    # Generate Button
    if st.button("📄 Generate Final Output"):
        with st.spinner("Generating Final Output..."):
            final_output = run_final_output_agent(
                understanding_json=st.session_state["understanding_json"],
                product_selection=st.session_state.get("product_result"),
                feedback=st.session_state.get("feedback_history", []) 
            )
            st.session_state["final_output"] = final_output

            st.session_state["report_is_saved"] = False

            _clear_analysis_state() 
            st.rerun()

# Display & Download Options
if "final_output" in st.session_state:
    st.markdown("### Final Report")
    
    # 1. Render on Screen
    report_text = st.session_state["final_output"]
    _render_report_with_image(st.session_state["final_output"])
    
    st.markdown("---")
    st.write("### 📥 Download Options")
    
    col1, col2, col3 = st.columns(3)
    
    # Option A: Download as Markdown (Original)
    with col1:
        st.download_button(
            label="Download as Markdown (.md)",
            data=st.session_state["final_output"],
            file_name="final_process_report.md",
            mime="text/markdown",
            on_click=_mark_as_saved
        )

    # Option B: Download as Word Doc (New)
    with col2:
        # Generate the DOCX file in memory on the fly
        docx_file = _generate_docx(
            st.session_state["final_output"], 
            image_path="process_map.png"
        )
        
        st.download_button(
            label="Download as Word Doc (.docx)",
            data=docx_file,
            file_name="final_process_report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            on_click=_mark_as_saved
        )

    # Option C: Save in App
    with col3:
        if st.button("💾 Save to App (Sidebar)", type="primary"):
            _save_report_to_app(report_text, "process_map.png")
            _mark_as_saved() # <--- Manually mark as saved
            st.rerun()
            st.success("Report saved in App sidebar!")

    if st.session_state.get("report_is_saved", True):
        st.caption("✅ Report saved. You can safely start a new analysis.")
    else:
        st.warning("⚠️ Report unsaved. Please save or download before starting a new analysis.")